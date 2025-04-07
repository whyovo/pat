import os
import tkinter as tk
import threading
from tkinter import messagebox
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
import re

from utils.api_utils import get_api_adapter, construct_prompt
from utils.app_manager import get_thread_safe_gui
from utils.thread_utils import ThreadSafeText


def generate_review(self):
    """启动文献综述生成过程"""
    if not self.excel_path.get() or not os.path.exists(self.excel_path.get()):
        self.status_bar["text"] = "错误：请先选择有效的Excel文件"
        messagebox.showerror("错误", "请先选择有效的Excel文件")
        return

    # 禁用界面按钮
    self.disable_analysis_buttons()

    # 更新状态
    self.status_bar["text"] = "准备开始综述生成..."
    self.output_text.delete(1.0, tk.END)  # 清空输出区域
    self.output_text.insert(tk.END, "=== 综述生成工具 ===\n\n", "header")
    self.output_text.insert(
        tk.END, "此功能将从Excel提取所有论文信息，生成完整的文献综述\n\n", "info"
    )
    self.output_text.tag_configure(
        "header", foreground="#5ba3e0", font=self.fonts["title"]
    )
    self.output_text.tag_configure("info", foreground=self.colors["fg"])
    self.output_text.see(tk.END)
    self.root.update_idletasks()

    try:
        api_url, api_key = self.get_api_info()

        if not api_key:
            self.status_bar["text"] = "错误：API密钥为空，无法进行分析"
            messagebox.showerror("错误", "API密钥为空，请在API设置中填写有效的密钥")
            self.enable_analysis_buttons()
            return

        # 显示正在处理的消息
        self.output_text.insert(tk.END, "正在启动综述生成，请稍候...\n", "processing")
        self.output_text.tag_configure(
            "processing", foreground="#4a8cca", font=("微软雅黑", 10)
        )
        self.output_text.see(tk.END)
        self.root.update_idletasks()

        threading.Thread(
            target=process_review_generation, args=(self, api_url, api_key), daemon=True
        ).start()
    except Exception as e:
        error_msg = f"启动综述生成失败: {str(e)}"
        self.status_bar["text"] = error_msg
        self.output_text.insert(tk.END, f"\n{error_msg}\n", "error")
        self.output_text.tag_configure("error", foreground="#cc0000")
        self.enable_analysis_buttons()
        import traceback

        print(f"综述生成启动错误: {str(e)}")
        print(traceback.format_exc())


def create_review_document(excel_path, review_content):
    """创建综述文档"""
    doc = Document()

    # 设置文档样式
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # 添加标题
    title = doc.add_heading("文献综述", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 处理markdown格式并添加内容
    lines = review_content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith("# "):
            # 一级标题
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            # 二级标题
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            # 三级标题
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            # 项目符号列表
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:])
        elif line.startswith("1. ") or line.startswith("1) "):
            # 编号列表
            p = doc.add_paragraph(style="List Number")
            p.add_run(line[3:])
        else:
            # 普通段落
            doc.add_paragraph(line)

        i += 1

    # 保存文档
    output_dir = os.path.dirname(excel_path)
    base_name = os.path.splitext(os.path.basename(excel_path))[0]
    output_path = os.path.join(output_dir, f"{base_name}_文献综述.docx")

    # 确保文件名不重复
    counter = 1
    original_path = output_path
    while os.path.exists(output_path):
        output_path = original_path.replace(".docx", f"_{counter}.docx")
        counter += 1

    doc.save(output_path)
    return output_path


def extract_all_excel_data(df):
    """从DataFrame提取所有列的数据"""
    all_data = {}

    # 提取每一列的数据
    for column in df.columns:
        # 去除NaN值并转换为字符串
        values = df[column].dropna().astype(str).tolist()
        # 过滤掉"未提供相关信息"和空字符串
        values = [
            v.strip() for v in values if v.strip() and v.strip() != "未提供相关信息"
        ]

        if values:  # 只添加非空列
            all_data[column] = values

    return all_data


def generate_review_content(adapter, all_data, terminate_check_fn=None):
    """使用AI生成综述内容"""
    # 构建系统提示词
    system_prompt = """你是一位学术论文综述撰写专家"""

    # 准备所有数据的字符串表示
    data_str = ""
    for column, values in all_data.items():
        data_str += f"## {column}:\n"
        for i, value in enumerate(values[:80], 1):
            data_str += f"{i}. {value}\n"
        if len(values) > 80:
            data_str += f"...（共{len(values)}条）\n"
        data_str += "\n"

    # 构建用户提示词
    user_prompt = f"""请基于以下论文信息，撰写一篇完整的文献综述：

{data_str}

要求至少包括以下部分：
1.研究背景与现状
2.国内外研究进展
3.主要研究方法与结论
4.研究空白与争议焦点
5.未来研究方向与展望
6.参考文献引用

请注意：
1.结构清晰，各部分主题鲜明，过渡自然
2.使用学术性语言，符合学术综述的写作规范
3.保持学术客观性，避免主观评价，不要添加任何未经过考证的内容
4.规范引用格式，在利用文献消息的时候必须在后面附上中括号括起来的数字，并在最后参考文献引用部分按照中国引用的规范格式列出。并且引用数字必须在整个文章中从1开始逐渐增大，不能同时出现两个相同的数字。
例子：Ghigliazza 等[11]展示了多足机器人在腿部着陆时如何选择固定的腿部角度才能导致渐进稳定的周期性步态。Deng 等[12]将 SLIP 归约模型推广到四足情况，分析其动力学相关特性，研究了不同刚度模型对于腿部动作空间、机体姿态角、弹性变化量的影响。最后写参考文献引用的时候：[11]Ghigliazza, R. S., Blankespoor, K., & Koditschek, D. E. (2006). A simple planar model of compliant legged locomotion. In Proceedings of the 2006 IEEE International Conference on Robotics and Automation (pp. 2390-2395). IEEE. [12]Deng, X., Avendano, A., & Ferrell, C. (2019). A simple planar model of compliant legged locomotion. In Proceedings of the 2019 IEEE International Conference on Robotics and Automation (pp. 2390-2395). IEEE.
"""

    try:
        # 修改为流式输出，将结果实时传递给UI
        print("开始流式生成综述内容...")
        # 传入正确的终止检查函数
        response = adapter.create_completion(
            user_prompt,
            system_prompt,
            stream=True,
            timeout=180,
            terminate_check_fn=terminate_check_fn,
        )

        result = ""
        for chunk in response:
            # 检查是否应该终止处理
            if terminate_check_fn and terminate_check_fn():
                print("综述生成任务被终止")
                yield "\n[任务已取消]\n"
                return "任务已取消"

            if hasattr(chunk, "choices") and len(chunk.choices) > 0:
                if hasattr(chunk.choices[0], "delta") and hasattr(
                    chunk.choices[0].delta, "content"
                ):
                    content = chunk.choices[0].delta.content
                    if content:
                        result += content
                        # 这里不再直接在此函数中更新UI，而是交由调用者处理
                        yield content  # 使用生成器模式返回每个片段

        return result
    except Exception as e:
        print(f"生成综述内容时出错: {str(e)}")
        yield f"\n生成综述内容时出错: {str(e)}\n"  # 也以生成器方式返回错误信息
        return f"生成综述内容时出错: {str(e)}"


def create_review_document(excel_path, review_content):
    """创建综述文档"""
    doc = Document()

    # 设置文档样式
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # 添加标题
    title = doc.add_heading("文献综述报告", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加文档说明
    doc.add_paragraph(
        "本综述基于Excel文件中的所有论文信息，自动生成的文献综述，涵盖研究背景、问题分析、方法评述、结果分析和未来展望等方面。"
    )

    # 添加基础信息
    doc.add_heading("综述基础信息", level=1)
    p = doc.add_paragraph()
    p.add_run("数据来源: ").bold = True
    p.add_run(os.path.basename(excel_path))

    # 添加生成日期
    import datetime

    p = doc.add_paragraph()
    p.add_run("生成日期: ").bold = True
    p.add_run(datetime.datetime.now().strftime("%Y-%m-%d"))

    # 添加分隔线
    doc.add_paragraph("_" * 80)

    # 添加综述内容
    for paragraph in review_content.split("\n"):
        if paragraph.strip():
            if paragraph.startswith("#"):
                # 处理Markdown标题
                level = paragraph.count("#")
                text = paragraph.strip("#").strip()
                doc.add_heading(text, level=min(level + 1, 9))
            else:
                p = doc.add_paragraph(paragraph)

    # 保存文档
    output_dir = os.path.dirname(excel_path)
    base_name = os.path.splitext(os.path.basename(excel_path))[0]
    output_path = os.path.join(output_dir, f"{base_name}_文献综述.docx")

    # 确保文件名不重复
    counter = 1
    original_path = output_path
    while os.path.exists(output_path):
        output_path = original_path.replace(".docx", f"_{counter}.docx")
        counter += 1

    doc.save(output_path)
    return output_path


def beautify_review_document(doc_path):
    """美化综述文档，增强视觉效果和专业性"""
    try:
        # 打开已有文档
        doc = Document(doc_path)

        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.0)

        # 设置文档样式
        styles = doc.styles

        # 更新标题样式
        for i in range(1, 5):  # 为1-4级标题设置样式
            if f"Heading {i}" in styles:
                heading_style = styles[f"Heading {i}"]
                heading_style.font.name = "微软雅黑"
                heading_style.font.size = Pt(16 - (i - 1) * 2)  # 随级别递减
                heading_style.font.bold = True
                heading_style.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)  # 蓝色
                heading_style.paragraph_format.keep_with_next = True

                # 为一级标题添加空间而不是边框
                if i == 1:
                    heading_style.paragraph_format.space_before = Pt(12)
                    heading_style.paragraph_format.space_after = Pt(8)
                    # 移除边框设置，有些文档格式不支持这个属性

        # 设置正文段落样式
        style = styles["Normal"]
        style.font.name = "微软雅黑"
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(10)  # 段后间距
        style.paragraph_format.first_line_indent = Pt(21)  # 首行缩进2个字符

        # 遍历所有段落应用格式
        for para in doc.paragraphs:
            # 跳过已有样式的标题段落
            if para.style.name.startswith("Heading"):
                continue

            # 对普通段落应用格式
            if len(para.text) > 0 and not para.text.startswith("_"):  # 跳过分隔线
                # 为引用添加特殊样式
                if "[" in para.text and "]" in para.text:
                    for run in para.runs:
                        if "[" in run.text and "]" in run.text:
                            citation_pattern = r"\[\d+\]"
                            import re

                            matches = re.findall(citation_pattern, run.text)
                            if matches:
                                # 强调引用标记
                                run.font.color.rgb = RGBColor(
                                    0x00, 0x66, 0x99
                                )  # 深蓝色
                                run.font.bold = True

        # 美化"参考文献"部分
        ref_section = None
        for i, para in enumerate(doc.paragraphs):
            if "参考文献" in para.text.lower() or "references" in para.text.lower():
                ref_section = i
                # 特殊样式处理参考文献标题
                para.style = styles["Heading 1"]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                break

        # 如果找到了参考文献部分，处理引用条目
        if ref_section is not None:
            for i in range(ref_section + 1, len(doc.paragraphs)):
                para = doc.paragraphs[i]
                if para.text.strip() and re.match(r"^\[\d+\]", para.text):
                    # 这是一个参考文献条目
                    para.paragraph_format.left_indent = Inches(0.3)
                    para.paragraph_format.hanging_indent = Inches(0.3)
                    para.paragraph_format.space_after = Pt(6)
                    para.paragraph_format.first_line_indent = Pt(0)  # 取消首行缩进

        # 设置封面标题更突出
        if len(doc.paragraphs) > 0:
            title_para = doc.paragraphs[0]
            if "文献综述" in title_para.text:
                title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in title_para.runs:
                    run.font.size = Pt(24)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x00, 0x33, 0x99)  # 深蓝色
                    # 如果需要模拟标题下的线条，可以为最后一个run添加下划线
                    if run == title_para.runs[-1]:
                        run.underline = True

        # 保存美化后的文档
        doc.save(doc_path)
        print(f"文档美化完成: {doc_path}")
        return True

    except Exception as e:
        import traceback

        print(f"美化文档时出错: {e}")
        print(traceback.format_exc())
        return False


def process_review_generation(self, api_url, api_key):
    """执行综述生成过程"""
    tm = get_thread_safe_gui(self.root)
    # 修复这里：正确初始化ThreadSafeText，传入文本组件和root窗口
    out = ThreadSafeText(self.output_text, self.root)

    try:
        print("开始综述生成过程")  # 添加调试信息

        # 重置取消标志，确保不会误判为已取消
        self.cancel_analysis_requested = False
        if hasattr(self, "terminate_all_tasks"):
            self.terminate_all_tasks = False

        # 获取API适配器
        out.insert(tk.END, "正在初始化API连接...\n", "step")
        # 使用更安全的方法设置标签
        tm.add_task(self.output_text.tag_configure, "step", foreground="#6aa84f")

        model_name = getattr(
            self, "api_model_var", tk.StringVar(value="deepseek-chat")
        ).get()

        print(f"创建API适配器，模型: {model_name}, URL: {api_url}")
        # 修复API参数顺序, 确保URL和API Key正确传递
        adapter = get_api_adapter(api_url, api_key, model_name)

        # 读取Excel文件
        excel_path = self.excel_path.get()
        out.insert(
            tk.END, f"正在分析Excel文件: {os.path.basename(excel_path)}\n", "process"
        )
        # 使用更安全的方法设置标签
        tm.add_task(self.output_text.tag_configure, "process", foreground="#4a8cca")
        df = pd.read_excel(excel_path)

        if df.empty:
            raise ValueError("Excel文件为空，无内容可供分析")

        # 检查行数
        row_count = len(df)
        if row_count == 0:
            raise ValueError("Excel文件中没有有效数据行")

        out.insert(
            tk.END,
            f"找到 {row_count} 行有效数据，包含 {len(df.columns)} 个字段\n\n",
            "info",
        )
        # 使用更安全的方法设置标签
        tm.add_task(self.output_text.tag_configure, "info", foreground="#6aa84f")

        # 提取Excel中的所有信息
        tm.add_task(self.status_bar.configure, text="正在提取Excel中的所有信息...")

        # 准备数据，将每一列的内容整理为字典
        all_data = extract_all_excel_data(df)

        tm.add_task(self.status_bar.configure, text="正在生成综述内容...")

        # 生成综述 - 修改为实时显示结果
        out.insert(tk.END, "\n开始生成综述，实时输出内容：\n", "subtitle")
        # 使用更安全的方法设置标签
        tm.add_task(
            self.output_text.tag_configure,
            "subtitle",
            foreground="#5ba3e0",
            font=("微软雅黑", 11, "bold"),
        )

        # 创建临时变量保存完整内容
        review_content = ""

        # 创建终止检查函数的闭包，并传递给generate_review_content
        terminate_check = (
            lambda: hasattr(self, "terminate_all_tasks") and self.terminate_all_tasks
        )

        # 使用生成器逐步获取并显示结果 - 并传递终止检查函数
        for content_chunk in generate_review_content(
            adapter, all_data, terminate_check
        ):
            # 检查是否应该终止处理 - 在这里直接检查
            if hasattr(self, "terminate_all_tasks") and self.terminate_all_tasks:
                out.insert(tk.END, "\n任务已取消\n", "error")
                tm.add_task(self.enable_analysis_buttons)
                return

            review_content += content_chunk
            out.insert(tk.END, content_chunk, "content")
            # 使用更安全的方法设置标签
            tm.add_task(self.output_text.tag_configure, "content", foreground="#e0e0e0")

        # 修改此处：只检查一种取消标志，防止冲突
        # 检查是否已经被取消（双重检查）
        if hasattr(self, "terminate_all_tasks") and self.terminate_all_tasks:
            out.insert(tk.END, "\n任务已取消\n", "error")
            tm.add_task(self.enable_analysis_buttons)
            return

        # 检查生成的内容是否完整
        if not check_review_completeness(review_content):
            out.insert(tk.END, "\n内容不完整，继续生成...\n", "warning")
            tm.add_task(self.output_text.tag_configure, "warning", foreground="#e69138")

            # 继续生成缺失的部分
            for content_chunk in generate_review_continuation(
                adapter, review_content, all_data, terminate_check
            ):
                # 检查是否应该终止处理 - 在这里直接检查
                if hasattr(self, "terminate_all_tasks") and self.terminate_all_tasks:
                    out.insert(tk.END, "\n任务已取消\n", "error")
                    tm.add_task(self.enable_analysis_buttons)
                    return

                review_content += content_chunk
                out.insert(tk.END, content_chunk, "content")
                # 使用更安全的方法设置标签
                tm.add_task(
                    self.output_text.tag_configure, "content", foreground="#e0e0e0"
                )

        # 更新进度
        tm.add_task(self.status_bar.configure, text="正在创建结果报告...")

        # 创建Word文档
        output_path = create_review_document(excel_path, review_content)

        # 添加显式日志，帮助调试文件保存
        print(f"综述文档已保存至: {output_path}")

        # 添加美化步骤
        tm.add_task(self.status_bar.configure, text="正在美化文档格式...")
        out.insert(tk.END, "\n正在对文档进行美化处理...\n", "process")

        try:
            # 美化文档
            beautify_review_document(output_path)
            out.insert(tk.END, "文档美化完成！\n", "success")
            # 使用更安全的方法设置标签
            tm.add_task(
                self.output_text.tag_configure,
                "success",
                foreground="#4caf50",
                font=("微软雅黑", 10, "bold"),
            )
        except Exception as e:
            out.insert(
                tk.END, f"文档美化过程中出错: {str(e)}，将使用基本格式\n", "warning"
            )
            # 使用更安全的方法设置标签
            tm.add_task(self.output_text.tag_configure, "warning", foreground="#e69138")

        # 更新进度并显示结果
        tm.add_task(
            self.status_bar.configure,
            text=f"综述生成完成！结果已保存至: {os.path.basename(output_path)}",
        )
        out.insert(
            tk.END,
            f"\n综述生成完成！结果已保存至: {os.path.basename(output_path)}\n",
            "note",
        )
        # 使用更安全的方法设置标签
        tm.add_task(
            self.output_text.tag_configure,
            "note",
            foreground="#999999",
            font=("微软雅黑", 9, "italic"),
        )

    except Exception as e:
        # 添加错误处理
        import traceback

        error_details = traceback.format_exc()
        print(f"综述生成错误: {e}")
        print(error_details)
        try:
            tm.add_task(
                self.output_text.insert,
                tk.END,
                f"\n综述生成过程中出错: {str(e)}\n",
                "error",
            )
            tm.add_task(
                self.output_text.tag_configure,
                "error",
                foreground="#cc0000",
                font=("微软雅黑", 10, "bold"),
            )
        except Exception as inner_e:
            print(f"错误处理过程中发生异常: {inner_e}")
        finally:
            tm.add_task(self.status_bar.configure, text=f"综述生成失败: {str(e)}")

    finally:
        print("综述生成过程完成")  # 添加调试信息
        tm.add_task(self.enable_analysis_buttons)


def check_review_completeness(content):
    """检查综述内容是否完整"""
    # 检测是否包含了所有必要章节
    required_sections = [
        "研究背景",
        "背景",
        "引言",
        "简介",  # 背景章节
        "研究进展",
        "文献综述",  # 国内外研究进展
        "研究方法",
        "方法",  # 方法章节
        "研究空白",
        "争议",
        "问题",  # 研究空白与争议
        "未来",
        "展望",
        "方向",  # 未来研究方向
        "参考文献",
        "引用",
        "References",  # 参考文献
    ]

    # 检查是否有参考文献格式的内容 [数字]
    has_citations = bool(re.search(r"\[\d+\]", content))

    # 检查是否有足够多的章节
    section_count = 0
    for section in required_sections:
        if section in content:
            section_count += 1

    # 检查总字符数是否足够多（至少5000字）
    enough_content = len(content) >= 5000

    # 检查是否包含"总结"或"结论"等结尾章节词语
    has_conclusion = any(word in content for word in ["总结", "结论", "小结", "总结性"])

    # 综合判断：文章足够长且有引用，或者包含足够多的章节且有结论
    return (enough_content and has_citations) or (section_count >= 4 and has_conclusion)


def identify_missing_sections(content):
    """识别综述中缺失的章节"""
    missing_sections = []

    # 检查必要的章节
    if not any(section in content for section in ["研究背景", "背景", "引言", "简介"]):
        missing_sections.append("研究背景")

    if not any(
        section in content for section in ["国内外研究进展", "研究进展", "文献综述"]
    ):
        missing_sections.append("国内外研究进展")

    if not any(section in content for section in ["研究方法", "方法", "研究手段"]):
        missing_sections.append("研究方法与结论")

    if not any(section in content for section in ["研究空白", "争议", "问题", "不足"]):
        missing_sections.append("研究空白与争议焦点")

    if not any(section in content for section in ["未来研究", "未来方向", "展望"]):
        missing_sections.append("未来研究方向与展望")

    # 检查是否包含参考文献
    if not any(
        section in content for section in ["参考文献", "引用文献", "References"]
    ):
        missing_sections.append("参考文献引用")
    elif not re.search(
        r"\[\d+\].*?[\u4e00-\u9fa5a-zA-Z]+.*?[\u4e00-\u9fa5a-zA-Z]+", content
    ):
        # 有参考文献标题但可能没有实际内容
        missing_sections.append("完整的参考文献列表")

    # 如果没有明显缺失的章节但内容较短，则建议扩展
    if not missing_sections and len(content) < 5000:
        missing_sections.append("内容扩展与丰富")

    return missing_sections


def generate_review_continuation(
    adapter, partial_content, all_data, terminate_check_fn=None
):
    """生成综述的后续部分"""
    # 分析已生成内容，确定哪些部分缺失
    missing_sections = identify_missing_sections(partial_content)

    # 构建系统提示词
    system_prompt = (
        """你是一位学术论文综述撰写专家，现在需要你继续完成一篇未完成的文献综述"""
    )

    # 准备部分数据作为上下文
    context_data = ""
    for column, values in all_data.items():
        if "标题" in column or "研究" in column or "结论" in column:
            context_data += f"## {column}:\n"
            for i, value in enumerate(values[:5], 1):  # 只取少量数据作为上下文
                context_data += f"{i}. {value}\n"
            context_data += "\n"

    # 获取已生成内容的最后1000个字符作为上下文
    last_part = (
        partial_content[-1000:] if len(partial_content) > 1000 else partial_content
    )

    # 构建用户提示词
    user_prompt = f"""我有一篇未完成的文献综述，请帮我继续完成。

已生成的内容结尾部分:
```
{last_part}
```

需要补充的部分包括但不限于：{', '.join(missing_sections)}

请基于以上内容和以下数据，继续撰写文献综述：

{context_data}

请注意：
1. 结构清晰，各部分主题鲜明，过渡自然
2. 使用学术性语言，符合学术综述的写作规范
3. 保持学术客观性，避免主观评价，不要添加任何未经过考证的内容
4. 规范引用格式，在利用文献消息的时候必须在后面附上中括号括起来的数字，并在最后参考文献引用部分按照中国引用的规范格式列出。并且引用数字必须在整个文章中从1开始逐渐增大，不能同时出现两个相同的数字。
"""

    try:
        # 修改为流式输出，将结果实时传递给UI
        print("开始流式生成综述内容...")
        # 传入正确的终止检查函数
        response = adapter.create_completion(
            user_prompt,
            system_prompt,
            stream=True,
            timeout=180,
            terminate_check_fn=terminate_check_fn,
        )

        result = ""
        for chunk in response:
            # 检查是否应该终止处理
            if terminate_check_fn and terminate_check_fn():
                print("综述生成任务被终止")
                yield "\n[任务已取消]\n"
                return "任务已取消"

            if hasattr(chunk, "choices") and len(chunk.choices) > 0:
                if hasattr(chunk.choices[0], "delta") and hasattr(
                    chunk.choices[0].delta, "content"
                ):
                    content = chunk.choices[0].delta.content
                    if content:
                        result += content
                        # 这里不再直接在此函数中更新UI，而是交由调用者处理
                        yield content  # 使用生成器模式返回每个片段

        return result
    except Exception as e:
        print(f"生成综述内容时出错: {str(e)}")
        yield f"\n生成综述内容时出错: {str(e)}\n"  # 也以生成器方式返回错误信息
        return f"生成综述内容时出错: {str(e)}"
