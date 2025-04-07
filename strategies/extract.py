import os
import tkinter as tk
from tkinter import messagebox
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
import re
import threading
from docx.shared import Pt, Inches, RGBColor
from utils.api_utils import get_api_adapter, construct_prompt
from utils.app_manager import get_thread_safe_gui
from utils.thread_utils import ThreadSafeText


def extract_content(self):
    """启动内容提取分析过程"""
    if not self.excel_path.get() or not os.path.exists(self.excel_path.get()):
        self.status_bar["text"] = "错误：请先选择有效的Excel文件"
        messagebox.showerror("错误", "请先选择有效的Excel文件")
        return

    # 禁用界面按钮
    self.disable_analysis_buttons()

    # 更新状态
    self.status_bar["text"] = "准备开始内容提取分析..."
    self.output_text.delete(1.0, tk.END)  # 清空输出区域
    self.output_text.insert(tk.END, "=== 内容提取分析工具 ===\n\n", "header")
    self.output_text.insert(
        tk.END,
        "此功能将分析Excel，归纳出构成这些论文的最低维度，并给出创新建议\n\n",
        "info",
    )

    self.output_text.tag_configure(
        "header", foreground="#5ba3e0", font=self.fonts["title"]
    )
    self.output_text.tag_configure("info", foreground=self.colors["fg"])
    self.output_text.see(tk.END)
    self.root.update_idletasks()

    try:
        api_url, api_key = self.get_api_info()

        if not api_key:
            self.status_bar["text"] = "错误：API密钥为空，无法进行分析"
            messagebox.showerror("错误", "API密钥为空，请在API设置中填写有效的密钥")
            self.enable_analysis_buttons()
            return
        # 显示正在处理的消息
        self.output_text.insert(tk.END, "正在启动，请稍候...\n", "processing")
        self.output_text.tag_configure(
            "processing", foreground="#4a8cca", font=("微软雅黑", 10)
        )
        self.output_text.see(tk.END)
        self.root.update_idletasks()

        threading.Thread(
            target=process_content_extraction,
            args=(self, api_url, api_key),
            daemon=True,
        ).start()
    except Exception as e:
        error_msg = f"启动内容提取分析失败: {str(e)}"
        self.status_bar["text"] = error_msg
        self.output_text.insert(tk.END, f"\n{error_msg}\n", "error")
        self.output_text.tag_configure("error", foreground="#cc0000")
        self.enable_analysis_buttons()
        import traceback

        print(f"内容提取启动错误: {str(e)}")
        print(traceback.format_exc())


def create_analysis_report(excel_path, analysis_result):
    """创建分析报告文档，加入创新评估指标"""
    doc = Document()

    # 设置文档样式
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # 添加标题
    title = doc.add_heading("分析报告", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加简介
    doc.add_paragraph(
        "本报告从Excel中的论文基本信息中归纳出最低维度，确保维度不多于四个。对这些维度进行组合，并列出相应组合包含的论文。对每个组合进行打分，提供合理的评价依据。"
    )

    for paragraph in analysis_result.split("\n"):
        if paragraph.strip():
            if paragraph.startswith("#"):
                # 处理Markdown标题
                level = paragraph.count("#")
                text = paragraph.strip("#").strip()
                doc.add_heading(text, level=min(level + 1, 9))
            else:
                p = doc.add_paragraph(paragraph)

    # 保存文档
    output_dir = os.path.dirname(excel_path)
    base_name = os.path.splitext(os.path.basename(excel_path))[0]
    output_path = os.path.join(output_dir, f"{base_name}_创新点评估分析.docx")

    # 确保文件名不重复
    counter = 1
    original_path = output_path
    while os.path.exists(output_path):
        output_path = original_path.replace(".docx", f"_{counter}.docx")
        counter += 1

    doc.save(output_path)
    return output_path


def extract_all_excel_data(df):
    """从DataFrame提取所有列的数据"""
    all_data = {}

    # 提取每一列的数据
    for column in df.columns:
        # 去除NaN值并转换为字符串
        values = df[column].dropna().astype(str).tolist()
        # 过滤掉"未提供相关信息"和空字符串
        values = [
            v.strip() for v in values if v.strip() and v.strip() != "未提供相关信息"
        ]

        if values:  # 只添加非空列
            all_data[column] = values

    return all_data


def process_content_extraction(self, api_url, api_key):
    """执行内容提取分析过程"""
    tm = get_thread_safe_gui(self.root)
    # 修复这里：正确初始化ThreadSafeText，传入文本组件和root窗口
    out = ThreadSafeText(self.output_text, self.root)

    try:
        print("开始内容提取分析过程")  # 添加调试信息
        # 使用更安全的方法设置标签
        tm.add_task(self.output_text.tag_configure, "step", foreground="#6aa84f")

        # 重置取消标志，确保不会误判为已取消
        self.cancel_analysis_requested = False
        if hasattr(self, "terminate_all_tasks"):
            self.terminate_all_tasks = False

        model_name = getattr(
            self, "api_model_var", tk.StringVar(value="deepseek-chat")
        ).get()

        try:
            print(f"创建API适配器，模型: {model_name}, URL: {api_url}")
            # 修复API参数顺序, 确保URL和API Key正确传递
            adapter = get_api_adapter(api_url, api_key, model_name)
        except Exception as e:
            print(f"API适配器创建失败: {e}")  # 添加调试信息
            raise ValueError(f"API适配器初始化失败: {str(e)}")

        # 读取Excel文件
        excel_path = self.excel_path.get()
        out.insert(
            tk.END, f"正在分析Excel文件: {os.path.basename(excel_path)}\n", "process"
        )
        out.tag_configure("process", foreground="#4a8cca")

        try:
            df = pd.read_excel(excel_path)
            all_data = extract_all_excel_data(df)
            row_count = len(df)

            if row_count == 0:
                raise ValueError("Excel文件中没有有效的研究问题及创新点数据")

            out.insert(tk.END, f"找到 {row_count} 行有效数据\n\n", "info")
            tm.add_task(self.status_bar.configure, text="正在工作中...")

            out.insert(tk.END, "\n开始生成创新点评估，实时显示分析结果：\n", "subtitle")
            out.tag_configure(
                "subtitle", foreground="#5ba3e0", font=("微软雅黑", 11, "bold")
            )

            # 创建取消检查函数
            terminate_check = (
                lambda: hasattr(self, "terminate_all_tasks")
                and self.terminate_all_tasks
            )

            # 创建临时变量保存完整内容
            analysis_result = ""
            for content_chunk in analyze_innovation_method_connections(
                adapter, all_data, terminate_check
            ):
                # 检查是否应该终止处理 - 在这里直接检查
                if terminate_check():
                    out.insert(tk.END, "\n任务已取消\n", "error")
                    tm.add_task(self.enable_analysis_buttons)
                    return

                analysis_result += content_chunk
                out.insert(tk.END, content_chunk, "content")
                out.tag_configure("content", foreground="#e0e0e0")

            # 检查是否已经被取消（双重检查）
            # 修改此处：只检查一种取消标志，防止冲突
            if hasattr(self, "terminate_all_tasks") and self.terminate_all_tasks:
                out.insert(tk.END, "\n任务已取消\n", "error")
                tm.add_task(self.enable_analysis_buttons)
                return

            tm.add_task(self.status_bar.configure, text="正在创建结果报告...")

            # 创建Word文档
            output_path = create_analysis_report(excel_path, analysis_result)

            # 添加显式日志，帮助调试文件保存
            print(f"报告文件已保存至: {output_path}")

            tm.add_task(
                self.status_bar.configure,
                text=f"分析完成！结果已保存至: {os.path.basename(output_path)}",
            )

            # 不再显示分析摘要，因为我们已经实时显示了完整内容
            out.insert(tk.END, f"\n✓ 分析已完成并保存至: {os.path.basename(output_path)}\n", "success")
            out.tag_configure(
                "success", foreground="#4caf50", font=("微软雅黑", 10, "bold")
            )
            out.insert(tk.END, "\n(完整分析已保存到Word文档)\n", "note")
            out.tag_configure(
                "note", foreground="#999999", font=("微软雅黑", 9, "italic")
            )

        except Exception as e:
            raise ValueError(f"处理Excel文件时出错: {str(e)}")

    except Exception as e:
        import traceback

        error_details = traceback.format_exc()
        print(f"内容提取分析错误: {e}")  # 添加调试信息
        print(error_details)  # 添加详细错误堆栈
        try:
            # 添加额外的错误处理，防止异常级联
            tm.add_task(self.output_text.insert, tk.END, f"\n内容提取分析过程中出错: {str(e)}\n", "error")
            tm.add_task(self.output_text.tag_configure, "error", foreground="#cc0000", font=("微软雅黑", 10, "bold"))
            tm.add_task(self.output_text.insert, tk.END, f"错误详情: {e}\n", "error-detail")
            tm.add_task(self.output_text.tag_configure, "error-detail", foreground="#cc0000")
        except Exception as inner_e:
            print(f"错误处理过程中发生异常: {inner_e}")
        finally:
            tm.add_task(self.status_bar.configure, text=f"内容提取分析失败: {str(e)}")

    finally:
        print("内容提取分析过程完成")  # 添加调试信息
        tm.add_task(self.enable_analysis_buttons)


def analyze_innovation_method_connections(adapter, all_data, terminate_check_fn=None):
    """分析创新点与研究方法之间的联系，并提供创新评估"""
    # 准备所有数据的字符串表示
    data_str = ""
    for column, values in all_data.items():
        data_str += f"## {column}:\n"
        for i, value in enumerate(values[:80], 1):
            data_str += f"{i}. {value}\n"
        if len(values) > 80:
            data_str += f"...（共{len(values)}条）\n"
        data_str += "\n"

    # 构建系统提示词
    system_prompt = """你是一位学术研究分析专家和数据归纳大师"""

    # 构建用户提示词
    user_prompt = f"""
- Skills: 你拥有数据处理、信息归纳、逻辑分析、论文研究内容理解以及组合评价的关键能力，能够运用专业知识和技能，对论文进行多维度分析和组合打分。
- Goals: 
  1. 从Excel中的论文基本信息中归纳出最低维度，确保维度不多于四个，尽量只给出三个，且每个维度之间是正交的，不耦合不互相包含，例如，研究对象、研究环境、研究方法。
  2. 对每个维度给出具体的例子，要求全面，越多越好。例：研究对象:果蝇、蚂蚁、竹节虫、蟑、水黾、蠛类...。研究方法:高速摄像、3D运动捕捉、肌电图(EMG)、神经记录、力板测量...。研究环境：光滑/粗糙表面、斜坡、颗粒介质(沙地)、水面、狭窄隧道、垂直培壁。
  3. 对这些维度进行组合，其组合不能与任何一篇论文相同，例：蚂蚁+3D运动捕捉+...，并列出相应组合包含的论文名称和近似的论文，要求全面，至少需要把组合的那几个维度对应的论文指出，最少3篇。同时对每个组合进行创新性打分，并给出可行性评估。

文献excel数据如下：
{data_str}
"""

    # 修改为流式输出，逐步返回结果
    print("开始流式生成分析...")
    # 传入正确的终止检查函数
    response = adapter.create_completion(
        user_prompt,
        system_prompt,
        stream=True,
        timeout=180,
        terminate_check_fn=terminate_check_fn,
    )

    result = ""
    for chunk in response:
        # 检查是否应该终止处理
        if terminate_check_fn and terminate_check_fn():
            print("创新点分析任务被终止")
            yield "\n[任务已取消]\n"
            return

        if hasattr(chunk, "choices") and len(chunk.choices) > 0:
            if hasattr(chunk.choices[0], "delta") and hasattr(
                chunk.choices[0].delta, "content"
            ):
                content = chunk.choices[0].delta.content
                if content:
                    result += content
                    # 使用生成器模式返回每个片段
                    yield content

    return result


def create_analysis_report(excel_path, analysis_result):
    """创建分析报告文档，加入创新评估指标"""
    doc = Document()

    # 设置文档样式
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # 添加标题
    title = doc.add_heading("分析报告", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加简介
    doc.add_paragraph(
        "本报告从Excel中的论文基本信息中归纳出最低维度，确保维度不多于四个。对这些维度进行组合，并列出相应组合包含的论文。对每个组合进行打分，提供合理的评价依据。"
    )

    for paragraph in analysis_result.split("\n"):
        if paragraph.strip():
            if paragraph.startswith("#"):
                # 处理Markdown标题
                level = paragraph.count("#")
                text = paragraph.strip("#").strip()
                doc.add_heading(text, level=min(level + 1, 9))
            else:
                p = doc.add_paragraph(paragraph)

    # 保存文档
    output_dir = os.path.dirname(excel_path)
    base_name = os.path.splitext(os.path.basename(excel_path))[0]
    output_path = os.path.join(output_dir, f"{base_name}_创新点评估分析.docx")

    # 确保文件名不重复
    counter = 1
    original_path = output_path
    while os.path.exists(output_path):
        output_path = original_path.replace(".docx", f"_{counter}.docx")
        counter += 1

    doc.save(output_path)
    return output_path
